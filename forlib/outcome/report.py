import docx
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
import os
import errno


class MdExport:
    def __init__(self, name):
        try:
            if not (os.path.isdir('result')):
                os.makedirs(os.path.join('result'))
        except OSError as e:
            if e.errno != errno.EEXIST:
                print("fail to create folder")
                raise
        file = open('./result/'+name+'.md', 'w', encoding='utf-8')
        self.file = file
        self.file.write('# <center>Analysis Report\n\n')
        self.file.write('<div style = "text-align: right"> '+str(datetime.now())+' </div>\n\n\n')

    def add_text(self, contents):
        self.file.write('\n'+contents+'\n\n')

    def emphasis(self, contents):
        self.file.write('\n\n**'+str(contents)+'**\n\n')

    def add_table(self, data):
        self.file.write('\n\n')
        if len(data) == 0:
            print("There is no data")
            return -1
        self.file.write('|')
        try:
            for i in data[0].keys():
                self.file.write('<center>'+i+'|')
            self.file.write('\n')
            for i in data[0].keys():
                self.file.write('--------------------|')
        except:
            print('There are some characters that we can\'t decode')

        self.file.write('\n')
        for i in range(0, len(data)):
            result = list(data[i].values())
            for j in range(0, len(data[0].keys())):
                self.file.write('<center>'+str(result[j])+'|')
            self.file.write('\n')
        self.file.write('\n\n')

    def save(self):
        self.file.close()


class DocxExport:
    def __init__(self):
        try:
            if not (os.path.isdir('result')):
                os.makedirs(os.path.join('result'))
        except OSError as e:
            if e.errno != errno.EEXIST:
                print("fail to create folder")
                raise
        document = docx.Document()
        heading = document.add_heading('Analysis Report', 0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        date_info = document.add_paragraph(str(datetime.now()))
        date_info.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        contents_info = document.add_paragraph()
        contents_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        self.document = document

    def add_table(self, data):
        if len(data) == 0:
            print("There is no data")
            return -1
        table = self.document.add_table(rows=1, cols=len(data[0].keys()))
        table.style = 'Table Grid'
        table.rows[0].style = "borderColor:red;background-color:gray"

        col_list = list(data[0].keys())

        hdr_cells = table.rows[0].cells
        for i in range(0, len(data[0].keys())):
            try:
                hdr_cells[i].text = str(col_list[i])
            except:
                hdr_cells[i].text = 'cannot put data'

        for i in range(0, len(data)):
            row_cells = table.add_row().cells
            result = list(data[i].values())
            for j in range(0, len(data[0].keys())):
                try:
                    row_cells[j].text = str(result[j])
                except:
                    row_cells[j].text = 'cannot put data'
        self.document.add_paragraph()

    def table_by_json(self, data):
        if len(data) == 0:
            print("There is no data")
            return -1
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].style = "borderColor:red;background-color:gray"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Key'
        hdr_cells[1].text = 'Value'

        try:
            for i in range(0, len(data.keys())):
                row_cells = table.add_row().cells
                row_cells[0].text = str(list(data.keys())[i])
                row_cells[1].text = str(list(data.values())[i])
        except ValueError:
            pass
        self.document.add_paragraph()

    def add_img(self, path):
        self.document.add_picture(path, width=Cm(15), height=Cm(15))

    def add_text(self, text):
        self.document.add_paragraph(str(text))

    def save(self, name):
        self.document.save('./result/' + name + '.docx')

